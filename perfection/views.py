# ==============================================================================
#  Copyright (C) 2022 Sakuyark, Inc. All Rights Reserved                       =
#                                                                              =
#    @Time : 2022-8-2 17:11                                                    =
#    @Author : hanjin                                                          =
#    @Email : 2819469337@qq.com                                                =
#    @File : views.py                                                          =
#    @Program: website                                                         =
# ==============================================================================
import random

from django.db.models import QuerySet
from django.http import HttpResponse
from django.utils import timezone
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.pagination import PageNumberPagination
from rest_framework.response import Response

from images.models import Image
from perfection.models.base import PerfectionStudent
from perfection.serializers.base import PerfectionStudentCreateSerializer, PerfectionStudentSerializer
from .conf import settings
from .models.words import WordsPerfection
from .serializers.words import (
    WordsPerfectionFinishSerializer, WordsPerfectionRememberAndReviewSerializer, WordsPerfectionRememberSerializer,
    WordsPerfectionReviewSerializer,
    WordsPerfectionSerializer, WordsPerfectionUnrememberedSerializer,
)


def to_word(words, mode="review"):
    if isinstance(words, QuerySet):
        words = list(words)
    if mode == 'review':
        random.shuffle(words)
    fontsize = 12
    doc = Document()
    doc.styles['Normal'].font.name = 'Consolas'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    doc.styles['Normal'].font.size = Pt(fontsize)
    if mode == 'review':
        doc.sections[0]._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
    doc.sections[0].top_margin = Cm(1.27)
    doc.sections[0].bottom_margin = Cm(1.27)
    doc.sections[0].left_margin = Cm(1.27)
    doc.sections[0].right_margin = Cm(1.27)
    for index, word in enumerate(words):
        if mode == 'review':
            para = doc.add_paragraph()
            para.add_run('%d. ' % (index + 1))
            para.add_run(word.word.word + ' ')
            para.add_run("_" * 15)
            para.paragraph_format.line_spacing = 1.5
        elif mode == 'remember':
            para = doc.add_paragraph()
            para.add_run('%d. ' % (index + 1))
            para.add_run(word.word.symbol + ' ')
            para.add_run(word.word.word + ' ').bold = True
            para.add_run(word.word.chinese)
            para.paragraph_format.line_spacing = 1.2
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    if mode == 'review':
        # 加入答案
        doc.add_section()
        doc.sections[1]._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
        for index, word in enumerate(words):
            para = doc.add_paragraph()
            para.add_run('%d. ' % (index + 1))
            para.add_run(word.word.symbol + ' ')
            para.add_run(word.word.word + ' ').bold = True
            para.add_run(word.word.chinese)
            para.paragraph_format.line_spacing = 1.2
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    return doc


# 生成word文档，该函数并未封装
def to_word_response(words, created, mode="review"):
    doc = to_word(words, mode)
    response = HttpResponse(content_type="application/octet-stream")
    response["Content-Disposition"] = \
        f'''attachment; filename={created.__format__("%Y-%m-%d-remember")}.docx'''.encode()
    response["Access-Control-Expose-Headers"] = "Content-Disposition"
    doc.save(response)
    return response


class PerfectionStudentViewSet(viewsets.ModelViewSet):
    serializer_class = PerfectionStudentSerializer
    queryset = PerfectionStudent.objects.all()
    permission_classes = settings.PERMISSIONS.student
    lookup_field = 'id'

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()
        if self.action == "list" and user.admin == 0:
            queryset = queryset.filter(pk=user.perfection.pk)
        return queryset

    def get_serializer_class(self):
        if self.action == 'create':
            return PerfectionStudentCreateSerializer
        return super().get_serializer_class()

    def get_permissions(self):
        if self.action == 'create':
            self.permission_classes = settings.PERMISSIONS.student_create
        elif self.action == "me":
            self.permission_classes = settings.PERMISSIONS.student
        return super().get_permissions()

    def get_instance(self):
        return self.request.user

    def create(self, request, *args, **kwargs):
        # 测试通过-正常
        user = self.get_instance()
        if getattr(user, 'perfection', None) is not None:
            return Response(status=status.HTTP_400_BAD_REQUEST, data={'perfection': ['已开通打卡功能']})
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        perfection = serializer.save()
        perfection.user = user
        perfection.save()
        return Response(status=status.HTTP_200_OK)

    @action(methods=['get'], detail=False)
    def me(self, request, *args, **kwargs):
        _object = self.get_instance().perfection
        serializer = self.get_serializer(_object)
        return Response(serializer.data)


class WordsPagination(PageNumberPagination):
    # 默认的大小
    page_size = 20
    page_size_query_param = 'page_size'
    max_page_size = 30


class WordsPerfectionViewSet(viewsets.ModelViewSet):
    serializer_class = WordsPerfectionSerializer
    queryset = WordsPerfection.objects.all()
    permission_classes = settings.PERMISSIONS.words
    lookup_field = 'id'
    pagination_class = WordsPagination

    def get_queryset(self):
        user = self.request.user
        # if user.admin > 0 or user.perfection.role == 'teacher':
        #     queryset = super().get_queryset()
        # else:
        queryset = user.perfection.words.all()
        return queryset

    def get_serializer_class(self):
        if self.action == 'remember' or self.action == 'remember_file':
            return WordsPerfectionRememberSerializer
        elif self.action == 'review' or self.action == 'review_file':
            return WordsPerfectionReviewSerializer
        elif self.action == 'unremembered':
            return WordsPerfectionUnrememberedSerializer
        elif self.action == 'remember_review':
            return WordsPerfectionRememberAndReviewSerializer
        elif self.action == 'finish':
            return WordsPerfectionFinishSerializer
        return super().get_serializer_class()

    def get_permissions(self):
        if self.action == 'create':
            self.permission_classes = settings.PERMISSIONS.words_create
        elif (self.action == 'remember' or self.action == 'review'
              or self.action == 'remember_file' or self.action == 'review_file'
              or self.action == 'remember_review' or self.action == 'unremembered'):
            self.permission_classes = settings.PERMISSIONS.words
        elif self.action == 'finish':
            self.permission_classes = settings.PERMISSIONS.words_finish
        return super().get_permissions()

    def get_instance(self):
        return self.request.user

    @action(methods=['get', 'post'], detail=True)
    def remember(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance=instance)
        data = serializer.data
        if not request.query_params.get('detail'):
            data['remember'] = [x['word'] for x in data['remember']]
        return Response(data=data, status=status.HTTP_200_OK)

    @action(methods=['get', 'post'], detail=True)
    def unremembered(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance=instance)
        data = serializer.data
        if not request.query_params.get('detail'):
            data['unremembered'] = [x['word'] for x in data['unremembered']]
        return Response(data=data, status=status.HTTP_200_OK)

    @action(methods=['get', 'post'], detail=True)
    def review(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance=instance)
        data = serializer.data
        if not request.query_params.get('detail'):
            data['review'] = [x['word'] for x in data['review']]
            if request.query_params.get('simple'):
                data['review'] = [x['word'] for x in data['review']]
        return Response(data=data, status=status.HTTP_200_OK)

    @action(methods=['get'], detail=True)
    def remember_file(self, request, *args, **kwargs):
        instance = self.get_object()
        return to_word_response(instance.remember.all(), created=instance.created, mode='remember')

    @action(methods=['get'], detail=True)
    def review_file(self, request, *args, **kwargs):
        instance = self.get_object()
        return to_word_response(instance.remember.all(), created=instance.created, mode='review')

    @action(methods=['get'], detail=True)
    def remember_review(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance=instance)
        data = serializer.data
        if not request.query_params.get('detail'):
            data['remember'] = [x['word'] for x in data['remember']]
            data['review'] = [x['word'] for x in data['review']]
        return Response(data=data, status=status.HTTP_200_OK)

    @action(methods=['post'], detail=True)
    def finish(self, request, *args, **kwargs):
        _object = self.get_object()
        if _object.is_finished:
            # 伪造drf验证返回值
            return Response(status=status.HTTP_400_BAD_REQUEST, data={'non_field_errors': ['打卡已完成']})
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data
        errors = {}
        review = _object.review.all()
        if review.count() != len(data["review"].keys()):
            errors['review'] = "「打卡版」词量不匹配"
        if len(errors.keys()) > 0:
            return Response(status=status.HTTP_400_BAD_REQUEST, data=errors)
        # 查验传入字典的单词
        review_errors = set(data['review']) - set(review.values_list('word__word', flat=True))
        if len(review_errors) > 0:
            errors['review'] = f"{review_errors} 不是本次「打卡版」的单词"
        if len(errors.keys()) > 0:
            return Response(status=status.HTTP_400_BAD_REQUEST, data=errors)
        _object.update_words(
            review=data['review']
        )
        # _object.picture = data['picture']
        # 处理图片问题
        for pic in data['picture']:
            image = Image.objects.create(image=pic)
            _object.picture.add(image)
        #
        _object.is_finished = True
        _object.finished = timezone.now()
        _object.save()
        picture = []
        for pic in _object.picture.all():
            picture.append(request.build_absolute_uri(pic.image.url))
        return Response(
            status=status.HTTP_200_OK, data={
                "accuracy": _object.accuracy,
                "total": _object.total,
                "picture": picture,
                "unremembered_words": _object.unremembered_words.values_list('word__word', 'word__chinese')
            }
        )
